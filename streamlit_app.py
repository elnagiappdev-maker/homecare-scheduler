# -*- coding: utf-8 -*-
# app.py
# Smart Homecare Scheduler (final single-file)
# Branding: All Rights Reserved © Dr. Yousra Abdelatti shown on login page bottom and inside app footer
#
# Requirements (put in requirements.txt in same repo):
# streamlit, pandas, altair, openpyxl, python-docx, matplotlib

import streamlit as st
import pandas as pd
import sqlite3
from datetime import datetime, date, time as dtime, timedelta
from io import BytesIO
import hashlib
import os
import tempfile
import matplotlib.pyplot as plt
import altair as alt
from docx import Document
from docx.shared import Inches

# ---------------------------
# Configuration / Constants
# ---------------------------
DB_PATH = "homecare_scheduler.db"
APP_TITLE = "Smart Homecare Scheduler (24/7)"
RELAXING_BG = "#E8F6F3"
ACCENT = "#5DADE2"
SMALL_FOOTER_STYLE = "font-size:12px; color:purple; font-weight:bold;"  # use purple for All Rights Reserved

STAFF_ROLES = ["Specialist", "GP", "Nurse", "RT", "PT", "Care Giver"]

# Fixed fields for ordering and labels (used for placement anchor when adding custom fields)
FIXED_FIELDS = {
    "patients": [
        ("id", "Patient ID"),
        ("name", "Full name"),
        ("dob", "Date of birth"),
        ("gender", "Gender"),
        ("phone", "Phone"),
        ("email", "Email"),
        ("address", "Address"),
        ("emergency_contact", "Emergency contact"),
        ("insurance_provider", "Insurance provider"),
        ("insurance_number", "Insurance number"),
        ("allergies", "Allergies"),
        ("medications", "Current medications"),
        ("diagnosis", "Primary diagnosis"),
        ("equipment_required", "Equipment required"),
        ("mobility", "Mobility"),
        ("care_plan", "Care plan summary"),
        ("notes", "Notes / social history"),
    ],
    "staff": [
        ("id", "Staff ID"),
        ("name", "Full name"),
        ("role", "Role"),
        ("license_number", "License number"),
        ("specialties", "Specialties"),
        ("phone", "Phone"),
        ("email", "Email"),
        ("availability", "Availability"),
        ("notes", "Notes"),
    ],
    "schedule": [
        ("visit_id", "Visit ID"),
        ("patient_id", "Patient ID"),
        ("staff_id", "Staff ID"),
        ("date", "Date"),
        ("start_time", "Start"),
        ("end_time", "End"),
        ("visit_type", "Visit type"),
        ("duration_minutes", "Duration (min)"),
        ("priority", "Priority"),
        ("notes", "Notes"),
    ]
}

# ---------------------------
# Utilities
# ---------------------------
def get_db_connection():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def hash_pw(pw: str) -> str:
    return hashlib.sha256(pw.encode()).hexdigest()

def now_iso():
    return datetime.utcnow().isoformat()

# ---------------------------
# DB init & migration (adds extra_fields & extra_values if absent)
# ---------------------------
def init_db_and_migrate():
    conn = get_db_connection()
    cur = conn.cursor()

    # fixed tables (safe create)
    cur.execute('''
        CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            password_hash TEXT,
            role TEXT,
            created_at TEXT
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS patients (
            id TEXT PRIMARY KEY,
            name TEXT,
            dob TEXT,
            gender TEXT,
            phone TEXT,
            email TEXT,
            address TEXT,
            emergency_contact TEXT,
            insurance_provider TEXT,
            insurance_number TEXT,
            allergies TEXT,
            medications TEXT,
            diagnosis TEXT,
            equipment_required TEXT,
            mobility TEXT,
            care_plan TEXT,
            notes TEXT,
            created_by TEXT,
            created_at TEXT
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS staff (
            id TEXT PRIMARY KEY,
            name TEXT,
            role TEXT,
            license_number TEXT,
            specialties TEXT,
            phone TEXT,
            email TEXT,
            availability TEXT,
            notes TEXT,
            created_by TEXT,
            created_at TEXT
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS schedule (
            visit_id TEXT PRIMARY KEY,
            patient_id TEXT,
            staff_id TEXT,
            date TEXT,
            start_time TEXT,
            end_time TEXT,
            visit_type TEXT,
            duration_minutes INTEGER,
            priority TEXT,
            notes TEXT,
            recurring_rule TEXT,
            created_by TEXT,
            created_at TEXT
        )
    ''')

    # dynamic fields tables
    cur.execute('''
        CREATE TABLE IF NOT EXISTS extra_fields (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            entity TEXT,
            field_name TEXT,
            field_type TEXT,
            field_order REAL,
            options TEXT
        )
    ''')
    cur.execute('''
        CREATE TABLE IF NOT EXISTS extra_values (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            entity TEXT,
            record_id TEXT,
            field_id INTEGER,
            value TEXT
        )
    ''')

    # seed admin & doctor if no users
    cur.execute("SELECT COUNT(*) as c FROM users")
    row = cur.fetchone()
    if row is None or row["c"] == 0:
        cur.execute("INSERT OR REPLACE INTO users (username,password_hash,role,created_at) VALUES (?,?,?,?)",
                    ("admin", hash_pw("1234"), "admin", now_iso()))
        cur.execute("INSERT OR REPLACE INTO users (username,password_hash,role,created_at) VALUES (?,?,?,?)",
                    ("doctor", hash_pw("abcd"), "doctor", now_iso()))

    conn.commit()
    conn.close()

init_db_and_migrate()
conn = get_db_connection()

# ---------------------------
# Cached reads
# ---------------------------
@st.cache_data(show_spinner=False)
def read_table(name: str):
    return pd.read_sql_query(f"SELECT * FROM {name}", conn)

# ---------------------------
# UI helpers & CSS
# ---------------------------
def inject_css():
    st.markdown(f"""
    <style>
    .stApp {{
        background: linear-gradient(180deg, {RELAXING_BG} 0%, white 100%);
    }}
    .big-title {{
        font-size:28px;
        font-weight:700;
        color: #0b3d91;
        margin-bottom: 0.2rem;
    }}
    .login-bottom {{
        text-align:center;
        margin-top: 1rem;
    }}
    </style>
    """, unsafe_allow_html=True)

def make_visit_id():
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) as c FROM schedule")
    count = cur.fetchone()["c"] + 1
    return f"V{count:05d}"

def render_footer(show_all_rights=True):
    st.markdown("---")
    if show_all_rights:
        st.markdown(f"<div style='text-align:center;'><span style='{SMALL_FOOTER_STYLE}'>All Rights Reserved © Dr. Yousra Abdelatti</span></div>", unsafe_allow_html=True)

# ---------------------------
# Authentication & session
# ---------------------------
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.user = None
    st.session_state.user_role = None

def login_user(username, password):
    cur = conn.cursor()
    cur.execute("SELECT password_hash, role FROM users WHERE username = ?", (username,))
    row = cur.fetchone()
    if row and hash_pw(password) == row[0]:
        st.session_state.logged_in = True
        st.session_state.user = username
        st.session_state.user_role = row[1]
        return True
    return False

def logout_user():
    keys = list(st.session_state.keys())
    for k in keys:
        if k.startswith("theme"):
            continue
        try:
            del st.session_state[k]
        except Exception:
            pass
    st.session_state.logged_in = False
    st.session_state.user = None
    st.session_state.user_role = None

# ---------------------------
# Extra fields helpers (admin-managed)
# ---------------------------
def get_extra_fields(entity):
    cur = conn.cursor()
    cur.execute("SELECT * FROM extra_fields WHERE entity=? ORDER BY field_order ASC", (entity,))
    rows = cur.fetchall()
    return [dict(r) for r in rows]

def add_extra_field(entity, field_name, field_type, anchor_field=None, position="below", options_text=""):
    cur = conn.cursor()
    # fetch current custom fields
    cur.execute("SELECT id, field_order FROM extra_fields WHERE entity=? ORDER BY field_order ASC", (entity,))
    existing = cur.fetchall()

    # Build tokens with orders: fixed have orders spaced by 1000
    tokens = []
    step = 1000
    for i, (k, lbl) in enumerate(FIXED_FIELDS[entity]):
        tokens.append((f"fixed:{k}", (i+1)*step))
    for r in existing:
        tokens.append((f"custom:{r['id']}", r['field_order']))
    tokens = sorted(tokens, key=lambda x: x[1])

    anchor_order = None
    if anchor_field:
        # anchor_field can be like "fixed:dob" or "custom:3"
        token = anchor_field
        for t, o in tokens:
            if t == token:
                anchor_order = o
                break

    if anchor_order is None:
        new_order = tokens[-1][1] + step if tokens else step
    else:
        if position == "above":
            prev_orders = [o for t,o in tokens if o < anchor_order]
            prev = prev_orders[-1] if prev_orders else anchor_order - step
            new_order = (prev + anchor_order) / 2.0
        else:
            next_orders = [o for t,o in tokens if o > anchor_order]
            nxt = next_orders[0] if next_orders else anchor_order + step
            new_order = (anchor_order + nxt) / 2.0

    cur.execute("INSERT INTO extra_fields (entity, field_name, field_type, field_order, options) VALUES (?,?,?,?,?)",
                (entity, field_name, field_type, float(new_order), options_text))
    conn.commit()
    renormalize_field_orders(entity)

def renormalize_field_orders(entity):
    cur = conn.cursor()
    cur.execute("SELECT id FROM extra_fields WHERE entity=? ORDER BY field_order ASC", (entity,))
    rows = cur.fetchall()
    for idx, r in enumerate(rows):
        cur.execute("UPDATE extra_fields SET field_order=? WHERE id=?", (float(idx+1), r['id']))
    conn.commit()

def remove_extra_field(field_id):
    cur = conn.cursor()
    cur.execute("DELETE FROM extra_values WHERE field_id=?", (field_id,))
    cur.execute("DELETE FROM extra_fields WHERE id=?", (field_id,))
    conn.commit()

def set_extra_value(entity, record_id, field_id, value):
    cur = conn.cursor()
    cur.execute("SELECT id FROM extra_values WHERE entity=? AND record_id=? AND field_id=?", (entity, record_id, field_id))
    row = cur.fetchone()
    if row:
        cur.execute("UPDATE extra_values SET value=? WHERE id=?", (value, row['id']))
    else:
        cur.execute("INSERT INTO extra_values (entity, record_id, field_id, value) VALUES (?,?,?,?)",
                    (entity, record_id, field_id, value))
    conn.commit()

def get_extra_value(entity, record_id, field_id):
    cur = conn.cursor()
    cur.execute("SELECT value FROM extra_values WHERE entity=? AND record_id=? AND field_id=?", (entity, record_id, field_id))
    row = cur.fetchone()
    return row['value'] if row else None

# ---------------------------
# Exports & Word/Charts helpers
# ---------------------------
def to_excel_bytes(dfs: dict):
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        for name, df in dfs.items():
            df.to_excel(writer, sheet_name=name[:31], index=False)
    output.seek(0)
    return output.getvalue()

def df_to_csv_bytes(df: pd.DataFrame):
    return df.to_csv(index=False).encode()

def _save_png_from_matplotlib(fig):
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    fig.savefig(tmp.name, bbox_inches='tight')
    tmp.close()
    with open(tmp.name, "rb") as f:
        data = f.read()
    try:
        os.unlink(tmp.name)
    except Exception:
        pass
    return data

def create_word_report(patients_df, staff_df, schedule_df, charts=None):
    doc = Document()
    doc.add_heading(APP_TITLE, level=1)
    doc.add_paragraph("Report generated: " + datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"))
    def add_df_section(title, df):
        doc.add_heading(title, level=2)
        if df is None or df.empty:
            doc.add_paragraph("No data.")
            return
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr = table.rows[0].cells
        for i, c in enumerate(df.columns):
            hdr[i].text = str(c)
        for _, r in df.iterrows():
            row_cells = table.add_row().cells
            for i, c in enumerate(df.columns):
                val = r[c]
                row_cells[i].text = "" if pd.isna(val) else str(val)
    add_df_section("Patients", patients_df)
    add_df_section("Staff", staff_df)
    add_df_section("Schedule", schedule_df)
    if charts:
        for title, png in charts.items():
            doc.add_page_break()
            doc.add_heading(title, level=2)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                tmp.write(png)
                tmp.flush()
                doc.add_picture(tmp.name, width=Inches(6))
            try:
                os.unlink(tmp.name)
            except Exception:
                pass
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f.getvalue()

# ---------------------------
# App layout & logic
# ---------------------------
st.set_page_config(page_title=APP_TITLE, layout="wide", initial_sidebar_state="expanded")
inject_css = lambda: None  # placeholder to satisfy later call; we'll inject CSS below
# actually inject CSS by calling the function body above
def _inject_css_real():
    st.markdown(f"""
    <style>
    .stApp {{
        background: linear-gradient(180deg, {RELAXING_BG} 0%, white 100%);
    }}
    .big-title {{
        font-size:28px;
        font-weight:700;
        color: #0b3d91;
        margin-bottom: 0.2rem;
    }}
    .login-bottom {{
        text-align:center;
        margin-top: 1rem;
    }}
    </style>
    """, unsafe_allow_html=True)
_inject_css_real()

# --- Login page (All Rights Reserved shown at bottom under login controls) ---
if not st.session_state.logged_in:
    st.markdown('<div class="big-title">Smart Homecare Scheduler Login</div>', unsafe_allow_html=True)
    col1, col2 = st.columns([1,1])
    with col1:
        username = st.text_input("Username", key="login_user")
        password = st.text_input("Password", type="password", key="login_pw")
        if st.button("Login"):
            if login_user(username, password):
                st.success(f"Welcome back, {st.session_state.user} ({st.session_state.user_role})")
                st.experimental_rerun()
            else:
                st.error("Invalid credentials")
    with col2:
        st.write("Demo accounts: admin / 1234  •  doctor / abcd")
        st.write("If you don't have an account ask the admin to create one (Settings > Create user).")
    # All Rights Reserved placed BELOW login boxes
    st.markdown("<div class='login-bottom'><span style='font-weight:bold; color:purple;'>All Rights Reserved © Dr. Yousra Abdelatti</span></div>", unsafe_allow_html=True)
    st.stop()

# --- Main app ---
st.sidebar.image("https://upload.wikimedia.org/wikipedia/commons/8/88/Patient_icon.svg", width=96)
st.sidebar.title("Menu")
menu = ["Dashboard","Patients","Staff","Schedule","Analytics","Emergency","Settings","Export & Backup","Logout"]
choice = st.sidebar.selectbox("Go to", menu)

st.markdown(f"<div class='big-title'>{APP_TITLE}</div>", unsafe_allow_html=True)

# ---------- Dashboard ----------
if choice == "Dashboard":
    st.subheader("Dashboard")
    patients_df = read_table("patients")
    staff_df = read_table("staff")
    schedule_df = read_table("schedule")

    c1, c2, c3 = st.columns(3)
    c1.metric("Patients", len(patients_df))
    c2.metric("Staff", len(staff_df))
    c3.metric("Scheduled Visits", len(schedule_df))

    st.markdown("---")
    st.write("Upcoming visits (next 30 days):")
    if len(schedule_df) > 0:
        schedule_df['date_dt'] = pd.to_datetime(schedule_df['date'], errors='coerce')
        upcoming = schedule_df[(schedule_df['date_dt'] >= pd.Timestamp(date.today())) & (schedule_df['date_dt'] <= pd.Timestamp(date.today()+timedelta(days=30)))]
        upcoming = upcoming.sort_values(['date','start_time']).head(100)
        st.dataframe(upcoming[['visit_id','patient_id','staff_id','date','start_time','end_time','visit_type','priority']])
    else:
        st.info("No visits scheduled yet.")

    # quick analytics previews
    st.markdown("### Quick analytics")
    col1, col2 = st.columns(2)
    with col1:
        if not patients_df.empty:
            dfp = patients_df.copy()
            dfp['dob_dt'] = pd.to_datetime(dfp['dob'], errors='coerce')
            dfp['age'] = ((pd.Timestamp(date.today()) - dfp['dob_dt']).dt.days // 365).fillna(0).astype(int)
            age_bins = pd.cut(dfp['age'], bins=[-1,0,1,5,12,18,40,65,200],
                              labels=["<1","1-5","6-12","13-18","19-40","41-65","66-200"])
            age_count = age_bins.value_counts().sort_index().reset_index()
            age_count.columns = ['age_group','count']
            st.altair_chart((alt.Chart(age_count).mark_bar(color=ACCENT).encode(x='age_group', y='count')).properties(height=240), use_container_width=True)
        else:
            st.info("Add patients to see age distribution.")
    with col2:
        if not schedule_df.empty:
            vtypes = schedule_df['visit_type'].fillna("Unknown").value_counts().reset_index()
            vtypes.columns = ['visit_type','count']
            st.altair_chart((alt.Chart(vtypes).mark_arc().encode(theta='count', color='visit_type')).properties(height=240), use_container_width=True)
        else:
            st.info("No visits to show distribution.")
    render_footer(show_all_rights=True)

# ---------- Patients ----------
elif choice == "Patients":
    st.subheader("Manage Patients")
    patients_df = read_table("patients")

    # helper to construct ordered render list combining fixed + custom fields
    def build_render_list(entity):
        fixed = FIXED_FIELDS[entity]
        tokens = []
        base_step = 1000
        for idx, (key, label) in enumerate(fixed):
            tokens.append({"type":"fixed", "key": key, "label": label, "order": (idx+1)*base_step})
        for cf in get_extra_fields(entity):
            tokens.append({"type":"custom", "key": f"custom_{cf['id']}", "label": cf['field_name'], "order": cf['field_order'], "meta": cf})
        tokens = sorted(tokens, key=lambda x: x['order'])
        return tokens

    # Add patient with dynamic custom fields
    with st.expander("Add new patient (full details and custom fields)", expanded=True):
        with st.form("add_patient_full", clear_on_submit=True):
            render_list = build_render_list("patients")
            fixed_vals = {}
            custom_vals = {}
            for item in render_list:
                if item['type'] == 'fixed':
                    k = item['key']
                    label = item['label']
                    if k == "id":
                        fixed_vals['id'] = st.text_input(label, key=f"new_pat_{k}")
                    elif k == "name":
                        fixed_vals['name'] = st.text_input(label, key=f"new_pat_{k}")
                    elif k == "dob":
                        fixed_vals['dob'] = st.date_input(label, value=date(1950,1,1), min_value=date(1900,1,1), key=f"new_pat_{k}")
                    elif k == "gender":
                        fixed_vals['gender'] = st.selectbox(label, ["Female","Male","Other","Prefer not to say"], key=f"new_pat_{k}")
                    elif k == "phone":
                        fixed_vals['phone'] = st.text_input(label, key=f"new_pat_{k}")
                    elif k == "email":
                        fixed_vals['email'] = st.text_input(label, key=f"new_pat_{k}")
                    elif k == "address":
                        fixed_vals['address'] = st.text_area(label, key=f"new_pat_{k}")
                    elif k == "emergency_contact":
                        fixed_vals['emergency_contact'] = st.text_input(label, key=f"new_pat_{k}")
                    elif k == "insurance_provider":
                        fixed_vals['insurance_provider'] = st.text_input(label, key=f"new_pat_{k}")
                    elif k == "insurance_number":
                        fixed_vals['insurance_number'] = st.text_input(label, key=f"new_pat_{k}")
                    elif k == "allergies":
                        fixed_vals['allergies'] = st.text_area(label, key=f"new_pat_{k}")
                    elif k == "medications":
                        fixed_vals['medications'] = st.text_area(label, key=f"new_pat_{k}")
                    elif k == "diagnosis":
                        fixed_vals['diagnosis'] = st.text_area(label, key=f"new_pat_{k}")
                    elif k == "equipment_required":
                        fixed_vals['equipment_required'] = st.text_area(label, key=f"new_pat_{k}")
                    elif k == "mobility":
                        fixed_vals['mobility'] = st.selectbox(label, ["Independent","Assisted","Wheelchair","Bedbound"], key=f"new_pat_{k}")
                    elif k == "care_plan":
                        fixed_vals['care_plan'] = st.text_area(label, key=f"new_pat_{k}")
                    elif k == "notes":
                        fixed_vals['notes'] = st.text_area(label, key=f"new_pat_{k}")
                    else:
                        fixed_vals[k] = st.text_input(label, key=f"new_pat_{k}")
                else:
                    cf = item['meta']
                    fid = cf['id']
                    ftype = cf['field_type']
                    label = cf['field_name']
                    options = cf['options'] or ""
                    key = f"new_pat_custom_{fid}"
                    if ftype == "text":
                        custom_vals[fid] = st.text_input(label, key=key)
                    elif ftype == "number":
                        custom_vals[fid] = st.number_input(label, key=key, format="%f")
                    elif ftype == "date":
                        custom_vals[fid] = st.date_input(label, key=key)
                    elif ftype == "dropdown":
                        opts = [o.strip() for o in options.split(",") if o.strip()]
                        custom_vals[fid] = st.selectbox(label, opts if opts else ["-- no options --"], key=key)
                    else:
                        custom_vals[fid] = st.text_input(label, key=key)
            submitted = st.form_submit_button("Add patient")
            if submitted:
                if not (fixed_vals.get('id') and fixed_vals.get('name')):
                    st.error("Patient ID and Name are required.")
                else:
                    cur = conn.cursor()
                    cur.execute("""INSERT OR REPLACE INTO patients
                                   (id,name,dob,gender,phone,email,address,emergency_contact,insurance_provider,insurance_number,allergies,medications,diagnosis,equipment_required,mobility,care_plan,notes,created_by,created_at)
                                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                                (fixed_vals.get('id'), fixed_vals.get('name'),
                                 fixed_vals.get('dob').isoformat() if isinstance(fixed_vals.get('dob'), date) else str(fixed_vals.get('dob')),
                                 fixed_vals.get('gender'), fixed_vals.get('phone'), fixed_vals.get('email'), fixed_vals.get('address'),
                                 fixed_vals.get('emergency_contact'), fixed_vals.get('insurance_provider'), fixed_vals.get('insurance_number'),
                                 fixed_vals.get('allergies'), fixed_vals.get('medications'), fixed_vals.get('diagnosis'),
                                 fixed_vals.get('equipment_required'), fixed_vals.get('mobility'), fixed_vals.get('care_plan'),
                                 fixed_vals.get('notes'), st.session_state.user, now_iso()))
                    conn.commit()
                    for fid, val in custom_vals.items():
                        if isinstance(val, date):
                            val = val.isoformat()
                        set_extra_value("patients", fixed_vals.get('id'), fid, "" if val is None else str(val))
                    st.success("Patient saved")
                    st.experimental_rerun()

    st.markdown("---")
    st.write("Existing patients:")
    pat_df = read_table("patients")
    st.dataframe(pat_df)

    # Edit / Delete patient (including custom fields)
    with st.expander("Edit / Delete patient"):
        if pat_df.empty:
            st.info("No patients")
        else:
            sel = st.selectbox("Select patient", pat_df['id'].tolist())
            row = pat_df[pat_df['id']==sel].iloc[0]
            can_edit = (st.session_state.user_role == "admin") or (row.get("created_by") == st.session_state.user)
            if not can_edit:
                st.warning("Only admin or the creator can edit/delete this patient.")
            render_list = build_render_list("patients")
            fixed_vals = {}
            custom_vals = {}
            for item in render_list:
                if item['type'] == 'fixed':
                    k = item['key']
                    label = item['label']
                    if k == "id":
                        st.write(f"Patient ID: **{row['id']}**")
                        fixed_vals['id'] = row['id']
                    elif k == "name":
                        fixed_vals['name'] = st.text_input(label, value=row['name'])
                    elif k == "dob":
                        dob_val = pd.to_datetime(row['dob'], errors='coerce')
                        dval = dob_val.date() if pd.notna(dob_val) else date(1950,1,1)
                        fixed_vals['dob'] = st.date_input(label, value=dval, min_value=date(1900,1,1))
                    elif k == "gender":
                        fixed_vals['gender'] = st.selectbox(label, ["Female","Male","Other","Prefer not to say"], index=0)
                    elif k == "phone":
                        fixed_vals['phone'] = st.text_input(label, value=row['phone'])
                    elif k == "email":
                        fixed_vals['email'] = st.text_input(label, value=row['email'])
                    elif k == "address":
                        fixed_vals['address'] = st.text_area(label, value=row['address'])
                    elif k == "emergency_contact":
                        fixed_vals['emergency_contact'] = st.text_input(label, value=row['emergency_contact'])
                    elif k == "insurance_provider":
                        fixed_vals['insurance_provider'] = st.text_input(label, value=row['insurance_provider'])
                    elif k == "insurance_number":
                        fixed_vals['insurance_number'] = st.text_input(label, value=row['insurance_number'])
                    elif k == "allergies":
                        fixed_vals['allergies'] = st.text_area(label, value=row['allergies'])
                    elif k == "medications":
                        fixed_vals['medications'] = st.text_area(label, value=row['medications'])
                    elif k == "diagnosis":
                        fixed_vals['diagnosis'] = st.text_area(label, value=row['diagnosis'])
                    elif k == "equipment_required":
                        fixed_vals['equipment_required'] = st.text_area(label, value=row['equipment_required'])
                    elif k == "mobility":
                        fixed_vals['mobility'] = st.selectbox(label, ["Independent","Assisted","Wheelchair","Bedbound"])
                    elif k == "care_plan":
                        fixed_vals['care_plan'] = st.text_area(label, value=row['care_plan'])
                    elif k == "notes":
                        fixed_vals['notes'] = st.text_area(label, value=row['notes'])
                    else:
                        fixed_vals[k] = st.text_input(label, value=row.get(k, ""))
                else:
                    cf = item['meta']
                    fid = cf['id']
                    ftype = cf['field_type']
                    label = cf['field_name']
                    options = cf['options'] or ""
                    curval = get_extra_value("patients", row['id'], fid)
                    key = f"edit_pat_custom_{fid}"
                    if ftype == "text":
                        custom_vals[fid] = st.text_input(label, value=curval if curval is not None else "", key=key)
                    elif ftype == "number":
                        try:
                            custom_vals[fid] = st.number_input(label, value=float(curval) if curval else 0.0, key=key, format="%f")
                        except Exception:
                            custom_vals[fid] = st.number_input(label, value=0.0, key=key, format="%f")
                    elif ftype == "date":
                        try:
                            d = pd.to_datetime(curval).date() if curval else date.today()
                        except Exception:
                            d = date.today()
                        custom_vals[fid] = st.date_input(label, value=d, key=key)
                    elif ftype == "dropdown":
                        opts = [o.strip() for o in options.split(",") if o.strip()]
                        custom_vals[fid] = st.selectbox(label, opts if opts else ["-- no options --"], index=0 if curval is None else (opts.index(curval) if curval in opts else 0), key=key)
                    else:
                        custom_vals[fid] = st.text_input(label, value=curval if curval is not None else "", key=key)
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Save changes") and can_edit:
                    cur = conn.cursor()
                    cur.execute("""UPDATE patients SET name=?,dob=?,gender=?,phone=?,email=?,address=?,emergency_contact=?,insurance_provider=?,insurance_number=?,allergies=?,medications=?,diagnosis=?,equipment_required=?,mobility=?,care_plan=?,notes=? WHERE id=?""",
                                (fixed_vals.get('name'), fixed_vals.get('dob').isoformat() if isinstance(fixed_vals.get('dob'), date) else str(fixed_vals.get('dob')),
                                 fixed_vals.get('gender'), fixed_vals.get('phone'), fixed_vals.get('email'), fixed_vals.get('address'), fixed_vals.get('emergency_contact'),
                                 fixed_vals.get('insurance_provider'), fixed_vals.get('insurance_number'), fixed_vals.get('allergies'), fixed_vals.get('medications'),
                                 fixed_vals.get('diagnosis'), fixed_vals.get('equipment_required'), fixed_vals.get('mobility'), fixed_vals.get('care_plan'), fixed_vals.get('notes'), sel))
                    conn.commit()
                    for fid, val in custom_vals.items():
                        if isinstance(val, date):
                            val = val.isoformat()
                        set_extra_value("patients", sel, fid, "" if val is None else str(val))
                    st.success("Patient updated")
                    st.experimental_rerun()
            with col2:
                if st.button("Delete patient") and can_edit:
                    cur = conn.cursor()
                    cur.execute("DELETE FROM patients WHERE id=?", (sel,))
                    cur.execute("DELETE FROM extra_values WHERE entity='patients' AND record_id=?", (sel,))
                    conn.commit()
                    st.success("Patient deleted")
                    st.experimental_rerun()

# ---------- Staff ----------
elif choice == "Staff":
    st.subheader("Manage Staff")
    staff_df = read_table("staff")

    def build_render_list_staff(entity):
        fixed = FIXED_FIELDS[entity]
        tokens = []
        base_step = 1000
        for idx, (key, label) in enumerate(fixed):
            tokens.append({"type":"fixed", "key": key, "label": label, "order": (idx+1)*base_step})
        for cf in get_extra_fields(entity):
            tokens.append({"type":"custom", "key": f"custom_{cf['id']}", "label": cf['field_name'], "order": cf['field_order'], "meta": cf})
        tokens = sorted(tokens, key=lambda x: x['order'])
        return tokens

    with st.expander("Add staff member"):
        with st.form("add_staff_full", clear_on_submit=True):
            render_list = build_render_list_staff("staff")
            fixed_vals = {}
            custom_vals = {}
            for item in render_list:
                if item['type'] == 'fixed':
                    k = item['key']
                    label = item['label']
                    if k == "id":
                        fixed_vals['id'] = st.text_input(label, key=f"new_staff_{k}")
                    elif k == "name":
                        fixed_vals['name'] = st.text_input(label, key=f"new_staff_{k}")
                    elif k == "role":
                        fixed_vals['role'] = st.selectbox(label, STAFF_ROLES, key=f"new_staff_{k}")
                    elif k == "license_number":
                        fixed_vals['license_number'] = st.text_input(label, key=f"new_staff_{k}")
                    elif k == "specialties":
                        fixed_vals['specialties'] = st.text_input(label, key=f"new_staff_{k}")
                    elif k == "phone":
                        fixed_vals['phone'] = st.text_input(label, key=f"new_staff_{k}")
                    elif k == "email":
                        fixed_vals['email'] = st.text_input(label, key=f"new_staff_{k}")
                    elif k == "availability":
                        fixed_vals['availability'] = st.text_area(label, key=f"new_staff_{k}")
                    elif k == "notes":
                        fixed_vals['notes'] = st.text_area(label, key=f"new_staff_{k}")
                    else:
                        fixed_vals[k] = st.text_input(label, key=f"new_staff_{k}")
                else:
                    cf = item['meta']
                    fid = cf['id']
                    ftype = cf['field_type']
                    label = cf['field_name']
                    options = cf['options'] or ""
                    key = f"new_staff_custom_{fid}"
                    if ftype == "text":
                        custom_vals[fid] = st.text_input(label, key=key)
                    elif ftype == "number":
                        custom_vals[fid] = st.number_input(label, key=key, format="%f")
                    elif ftype == "date":
                        custom_vals[fid] = st.date_input(label, key=key)
                    elif ftype == "dropdown":
                        opts = [o.strip() for o in options.split(",") if o.strip()]
                        custom_vals[fid] = st.selectbox(label, opts if opts else ["-- no options --"], key=key)
                    else:
                        custom_vals[fid] = st.text_input(label, key=key)
            submitted = st.form_submit_button("Add staff")
            if submitted:
                if not (fixed_vals.get('id') and fixed_vals.get('name')):
                    st.error("Staff ID and Name required.")
                else:
                    cur = conn.cursor()
                    cur.execute("""INSERT OR REPLACE INTO staff (id,name,role,license_number,specialties,phone,email,availability,notes,created_by,created_at)
                                   VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
                                (fixed_vals.get('id'), fixed_vals.get('name'), fixed_vals.get('role'), fixed_vals.get('license_number'),
                                 fixed_vals.get('specialties'), fixed_vals.get('phone'), fixed_vals.get('email'), fixed_vals.get('availability'),
                                 fixed_vals.get('notes'), st.session_state.user, now_iso()))
                    conn.commit()
                    for fid, val in custom_vals.items():
                        if isinstance(val, date):
                            val = val.isoformat()
                        set_extra_value("staff", fixed_vals.get('id'), fid, "" if val is None else str(val))
                    st.success("Staff saved")
                    st.experimental_rerun()

    st.markdown("---")
    st.dataframe(staff_df)

    with st.expander("Edit / Delete staff"):
        if staff_df.empty:
            st.info("No staff yet")
        else:
            sel = st.selectbox("Select staff", staff_df['id'].tolist())
            row = staff_df[staff_df['id']==sel].iloc[0]
            can_edit = (st.session_state.user_role == "admin") or (row.get("created_by") == st.session_state.user)
            if not can_edit:
                st.warning("Only admin or creator can edit/delete this staff record.")
            render_list = build_render_list_staff("staff")
            fixed_vals = {}
            custom_vals = {}
            for item in render_list:
                if item['type'] == 'fixed':
                    k = item['key']
                    label = item['label']
                    if k == "id":
                        st.write(f"Staff ID: **{row['id']}**")
                        fixed_vals['id'] = row['id']
                    elif k == "name":
                        fixed_vals['name'] = st.text_input(label, value=row['name'])
                    elif k == "role":
                        fixed_vals['role'] = st.selectbox(label, STAFF_ROLES, index=STAFF_ROLES.index(row['role']) if row['role'] in STAFF_ROLES else 0)
                    elif k == "license_number':
                        fixed_vals['license_number'] = st.text_input(label, value=row['license_number'])
                    elif k == "specialties":
                        fixed_vals['specialties'] = st.text_input(label, value=row['specialties'])
                    elif k == "phone":
                        fixed_vals['phone'] = st.text_input(label, value=row['phone'])
                    elif k == "email":
                        fixed_vals['email'] = st.text_input(label, value=row['email'])
                    elif k == "availability":
                        fixed_vals['availability'] = st.text_area(label, value=row['availability'])
                    elif k == "notes":
                        fixed_vals['notes'] = st.text_area(label, value=row['notes'])
                    else:
                        fixed_vals[k] = st.text_input(label, value=row.get(k, ""))
                else:
                    cf = item['meta']
                    fid = cf['id']
                    ftype = cf['field_type']
                    label = cf['field_name']
                    options = cf['options'] or ""
                    curval = get_extra_value("staff", row['id'], fid)
                    key = f"edit_staff_custom_{fid}"
                    if ftype == "text":
                        custom_vals[fid] = st.text_input(label, value=curval if curval is not None else "", key=key)
                    elif ftype == "number":
                        try:
                            custom_vals[fid] = st.number_input(label, value=float(curval) if curval else 0.0, key=key, format="%f")
                        except Exception:
                            custom_vals[fid] = st.number_input(label, value=0.0, key=key, format="%f")
                    elif ftype == "date":
                        try:
                            d = pd.to_datetime(curval).date() if curval else date.today()
                        except Exception:
                            d = date.today()
                        custom_vals[fid] = st.date_input(label, value=d, key=key)
                    elif ftype == "dropdown":
                        opts = [o.strip() for o in options.split(",") if o.strip()]
                        custom_vals[fid] = st.selectbox(label, opts if opts else ["-- no options --"], index=0 if curval is None else (opts.index(curval) if curval in opts else 0), key=key)
                    else:
                        custom_vals[fid] = st.text_input(label, value=curval if curval is not None else "", key=key)
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Save staff") and can_edit:
                    cur = conn.cursor()
                    cur.execute("""UPDATE staff SET name=?,role=?,license_number=?,specialties=?,phone=?,email=?,availability=?,notes=? WHERE id=?""",
                                (fixed_vals.get('name'), fixed_vals.get('role'), fixed_vals.get('license_number'), fixed_vals.get('specialties'), fixed_vals.get('phone'), fixed_vals.get('email'), fixed_vals.get('availability'), fixed_vals.get('notes'), sel))
                    conn.commit()
                    for fid, val in custom_vals.items():
                        if isinstance(val, date):
                            val = val.isoformat()
                        set_extra_value("staff", sel, fid, "" if val is None else str(val))
                    st.success("Saved")
                    st.experimental_rerun()
            with col2:
                if st.button("Delete staff") and can_edit:
                    cur = conn.cursor()
                    cur.execute("DELETE FROM staff WHERE id=?", (sel,))
                    cur.execute("DELETE FROM extra_values WHERE entity='staff' AND record_id=?", (sel,))
                    conn.commit()
                    st.success("Deleted")
                    st.experimental_rerun()
    render_footer(show_all_rights=True)

# ---------- Schedule ----------
elif choice == "Schedule":
    st.subheader("Scheduling & Visits")
    patients_df = read_table("patients")
    staff_df = read_table("staff")
    schedule_df = read_table("schedule")

    # Build render list for schedule (fixed + custom)
    render_list_schedule = []
    base_step = 1000
    for idx, (key, label) in enumerate(FIXED_FIELDS['schedule']):
        render_list_schedule.append({"type":"fixed","key":key,"label":label,"order":(idx+1)*base_step})
    for cf in get_extra_fields("schedule"):
        render_list_schedule.append({"type":"custom","key":f"custom_{cf['id']}","label":cf['field_name'],"order":cf['field_order'],"meta":cf})
    render_list_schedule = sorted(render_list_schedule, key=lambda x: x['order'])

    col1, col2 = st.columns([2,1])
    with col1:
        st.markdown("### Create visit")
        if patients_df.empty:
            st.warning("Add patients first")
        if staff_df.empty:
            st.warning("Add staff first")
        with st.form("create_visit_form", clear_on_submit=True):
            fixed_vals = {}
            custom_vals = {}
            for item in render_list_schedule:
                if item['type'] == 'fixed':
                    k = item['key']
                    if k == "visit_id":
                        pass
                    elif k == "patient_id":
                        fixed_vals['patient_id'] = st.selectbox("Patient", patients_df['id'].tolist() if len(patients_df)>0 else [], key="visit_patient")
                    elif k == "staff_id":
                        fixed_vals['staff_id'] = st.selectbox("Staff", staff_df['id'].tolist() if len(staff_df)>0 else [], key="visit_staff")
                    elif k == "date":
                        fixed_vals['date'] = st.date_input("Date", value=date.today(), key="visit_date")
                    elif k == "start_time":
                        fixed_vals['start_time'] = st.time_input("Start", value=dtime(9,0), key="visit_start")
                    elif k == "end_time":
                        fixed_vals['end_time'] = st.time_input("End", value=dtime(10,0), key="visit_end")
                    elif k == "visit_type":
                        fixed_vals['visit_type'] = st.selectbox("Visit type", ["Home visit","Telehealth","Wound care","Medication administration","Physiotherapy","Respiratory therapy","Assessment","Other"], key="visit_type")
                    elif k == "priority":
                        fixed_vals['priority'] = st.selectbox("Priority", ["Low","Normal","High","Critical"], key="visit_priority")
                    elif k == "notes":
                        fixed_vals['notes'] = st.text_area("Notes / visit plan", key="visit_notes")
                    else:
                        fixed_vals[k] = st.text_input(item['label'])
                else:
                    cf = item['meta']
                    fid = cf['id']
                    ftype = cf['field_type']
                    label = cf['field_name']
                    options = cf['options'] or ""
                    key = f"new_sched_custom_{fid}"
                    if ftype == "text":
                        custom_vals[fid] = st.text_input(label, key=key)
                    elif ftype == "number":
                        custom_vals[fid] = st.number_input(label, key=key, format="%f")
                    elif ftype == "date":
                        custom_vals[fid] = st.date_input(label, key=key)
                    elif ftype == "dropdown":
                        opts = [o.strip() for o in options.split(",") if o.strip()]
                        custom_vals[fid] = st.selectbox(label, opts if opts else ["-- no options --"], key=key)
                    else:
                        custom_vals[fid] = st.text_input(label, key=key)
            submitted = st.form_submit_button("Create visit")
            if submitted:
                if not fixed_vals.get('patient_id') or not fixed_vals.get('staff_id'):
                    st.error("Select patient and staff")
                else:
                    vid = make_visit_id()
                    duration = int((datetime.combine(date.today(), fixed_vals.get('end_time')) - datetime.combine(date.today(), fixed_vals.get('start_time'))).seconds / 60)
                    cur = conn.cursor()
                    cur.execute("""INSERT OR REPLACE INTO schedule (visit_id,patient_id,staff_id,date,start_time,end_time,visit_type,duration_minutes,priority,notes,created_by,created_at)
                                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?)""",
                                (vid, fixed_vals.get('patient_id'), fixed_vals.get('staff_id'),
                                 fixed_vals.get('date').isoformat() if isinstance(fixed_vals.get('date'), date) else str(fixed_vals.get('date')),
                                 fixed_vals.get('start_time').strftime("%H:%M") if fixed_vals.get('start_time') else "",
                                 fixed_vals.get('end_time').strftime("%H:%M") if fixed_vals.get('end_time') else "",
                                 fixed_vals.get('visit_type'), duration, fixed_vals.get('priority'), fixed_vals.get('notes'), st.session_state.user, now_iso()))
                    conn.commit()
                    for fid, val in custom_vals.items():
                        if isinstance(val, date):
                            val = val.isoformat()
                        set_extra_value("schedule", vid, fid, "" if val is None else str(val))
                    st.success(f"Visit {vid} created")
                    st.experimental_rerun()
    with col2:
        st.markdown("### View / Manage visits")
        if schedule_df.empty:
            st.info("No visits scheduled yet.")
        else:
            sel_visit = st.selectbox("Select visit", schedule_df['visit_id'].tolist())
            row = schedule_df[schedule_df['visit_id']==sel_visit].iloc[0]
            st.write(row.to_dict())
            can_edit = (st.session_state.user_role == "admin") or (row.get("created_by") == st.session_state.user)
            if can_edit:
                if st.button("Delete visit"):
                    cur = conn.cursor()
                    cur.execute("DELETE FROM schedule WHERE visit_id=?", (sel_visit,))
                    cur.execute("DELETE FROM extra_values WHERE entity='schedule' AND record_id=?", (sel_visit,))
                    conn.commit()
                    st.success("Visit deleted")
                    st.experimental_rerun()
            else:
                st.info("Only admin or creator can delete this visit.")
    render_footer(show_all_rights=True)

# ---------- Analytics ----------
elif choice == "Analytics":
    st.subheader("Analytics")
    patients_df = read_table("patients")
    schedule_df = read_table("schedule")

    st.markdown("### Patients by age group")
    if not patients_df.empty:
        dfp = patients_df.copy()
        dfp['dob_dt'] = pd.to_datetime(dfp['dob'], errors='coerce')
        dfp['age'] = ((pd.Timestamp(date.today()) - dfp['dob_dt']).dt.days // 365).fillna(0).astype(int)
        bins = pd.cut(dfp['age'], bins=[-1,0,1,5,12,18,40,65,200],
                      labels=["<1","1-5","6-12","13-18","19-40","41-65","66+"])
        age_count = bins.value_counts().sort_index().reset_index()
        age_count.columns = ['age_group','count']
        st.altair_chart(alt.Chart(age_count).mark_bar(color=ACCENT).encode(x='age_group', y='count'), use_container_width=True)
    else:
        st.info("Add patients to see analytics.")

    st.markdown("### Staff workload (visits per staff)")
    if not schedule_df.empty:
        w = schedule_df['staff_id'].value_counts().reset_index()
        w.columns = ['staff_id','visits']
        st.altair_chart(alt.Chart(w).mark_bar(color="#66c2a5").encode(x='staff_id', y='visits'), use_container_width=True)
    else:
        st.info("No schedule data")
    render_footer(show_all_rights=True)

# ---------- Emergency ----------
elif choice == "Emergency":
    st.subheader("Emergency")
    patients_df = read_table("patients")
    if len(patients_df):
        sel = st.selectbox("Patient", patients_df['id'].tolist())
        row = patients_df[patients_df['id']==sel].iloc[0]
        st.write(row.to_dict())
        if st.button("Call emergency contact (mock)"):
            st.info("Calling: " + str(row['emergency_contact']))
    else:
        st.info("No patients yet.")
    render_footer(show_all_rights=True)

# ---------- Settings & Custom Fields management ----------
elif choice == "Settings":
    st.subheader("Settings & User Management")
    st.write(f"Logged in as **{st.session_state.user}** ({st.session_state.user_role})")

    # Change own password
    with st.expander("Change your password", expanded=True):
        old = st.text_input("Current password", type="password", key="old_pw")
        new = st.text_input("New password", type="password", key="new_pw")
        new2 = st.text_input("Confirm new password", type="password", key="new_pw2")
        if st.button("Change password"):
            if not old or not new or new != new2:
                st.error("Ensure fields are filled and new passwords match.")
            else:
                cur = conn.cursor()
                cur.execute("SELECT password_hash FROM users WHERE username=?", (st.session_state.user,))
                row = cur.fetchone()
                if row and hash_pw(old) == row[0]:
                    cur.execute("UPDATE users SET password_hash=? WHERE username=?", (hash_pw(new), st.session_state.user))
                    conn.commit()
                    st.success("Password changed.")
                else:
                    st.error("Current password incorrect.")

    # Admin user management & custom fields
    if st.session_state.user_role == "admin":
        st.markdown("### Admin: User Management")
        users_df = read_table("users")
        st.dataframe(users_df[['username','role','created_at']] if not users_df.empty else pd.DataFrame())
        with st.expander("Create new user"):
            u_name = st.text_input("Username")
            u_role = st.selectbox("Role", ["admin","doctor","nurse","staff","other"])
            u_pw = st.text_input("Password", type="password")
            if st.button("Create user"):
                if not u_name or not u_pw:
                    st.error("Username and password required")
                else:
                    cur = conn.cursor()
                    cur.execute("INSERT OR REPLACE INTO users (username,password_hash,role,created_at) VALUES (?,?,?,?)",
                                (u_name, hash_pw(u_pw), u_role, now_iso()))
                    conn.commit()
                    st.success("User created")
                    st.experimental_rerun()
        with st.expander("Reset user password"):
            users = read_table("users")
            if not users.empty:
                sel = st.selectbox("Select user", users['username'].tolist())
                new_pw = st.text_input("New password for selected user", type="password", key="reset_pw")
                if st.button("Reset password"):
                    if not new_pw:
                        st.error("Enter a password")
                    else:
                        cur = conn.cursor()
                        cur.execute("UPDATE users SET password_hash=? WHERE username=?", (hash_pw(new_pw), sel))
                        conn.commit()
                        st.success("Password reset")
            else:
                st.info("No users found")

        # Manage Custom Fields
        st.markdown("### Manage Custom Fields (Admin only)")
        st.write("Create/remove/reorder/place fields for Patients, Staff, Schedule. These appear automatically for all users.")
        with st.form("add_custom_field", clear_on_submit=True):
            entity = st.selectbox("Entity", ["patients", "staff", "schedule"])
            field_name = st.text_input("Field label (e.g. Drug history)")
            field_type = st.selectbox("Field type", ["text","number","date","dropdown"])
            options_text = ""
            if field_type == "dropdown":
                options_text = st.text_area("Dropdown options (comma separated)")
            # anchor choices include fixed and existing custom
            anchor_choices = ["(end)"]
            for k, lab in FIXED_FIELDS[entity]:
                anchor_choices.append(f"fixed:{k}|{lab}")
            existing = get_extra_fields(entity)
            for cf in existing:
                anchor_choices.append(f"custom:{cf['id']}|{cf['field_name']}")
            anchor_choice = st.selectbox("Place relative to", options=anchor_choices, index=0)
            position = st.selectbox("Position", ["below","above"])
            submitted = st.form_submit_button("Create custom field")
            if submitted:
                if not field_name.strip():
                    st.error("Field label required")
                else:
                    anchor_field = None
                    if anchor_choice != "(end)":
                        if anchor_choice.startswith("fixed:"):
                            anchor_field = anchor_choice.split("|")[0]  # fixed:dob
                        elif anchor_choice.startswith("custom:"):
                            anchor_field = anchor_choice.split("|")[0]  # custom:3
                    add_extra_field(entity, field_name.strip(), field_type, anchor_field=anchor_field, position=position, options_text=options_text)
                    st.success("Custom field added")
                    st.experimental_rerun()

        st.markdown("Existing custom fields")
        cf_tbl = read_table("extra_fields")
        if not cf_tbl.empty:
            st.dataframe(cf_tbl[['id','entity','field_name','field_type','field_order','options']])
            with st.form("remove_custom_field"):
                rem_id = st.number_input("Enter ID to remove", min_value=1, step=1)
                if st.form_submit_button("Remove field"):
                    remove_extra_field(int(rem_id))
                    st.success("Removed")
                    st.experimental_rerun()
        else:
            st.info("No custom fields yet.")
    else:
        st.info("Only admin can manage custom fields and users.")

    render_footer(show_all_rights=True)

# ---------- Export & Backup ----------
elif choice == "Export & Backup":
    st.subheader("Export & Backup")
    patients_df = read_table("patients")
    staff_df = read_table("staff")
    schedule_df = read_table("schedule")

    st.markdown("### Export data")
    c1, c2, c3 = st.columns(3)
    with c1:
        csv_pat = df_to_csv_bytes(patients_df) if not patients_df.empty else b""
        st.download_button("Download Patients CSV", data=csv_pat, file_name="patients.csv", mime="text/csv")
        csv_staff = df_to_csv_bytes(staff_df) if not staff_df.empty else b""
        st.download_button("Download Staff CSV", data=csv_staff, file_name="staff.csv", mime="text/csv")
        csv_sched = df_to_csv_bytes(schedule_df) if not schedule_df.empty else b""
        st.download_button("Download Schedule CSV", data=csv_sched, file_name="schedule.csv", mime="text/csv")
    with c2:
        excel_bytes = to_excel_bytes({"patients":patients_df, "staff":staff_df, "schedule":schedule_df})
        st.download_button("Download Excel (all)", data=excel_bytes, file_name="homecare_data.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    with c3:
        # prepare charts for embedding into Word
        charts = {}
        try:
            if not patients_df.empty:
                dfp = patients_df.copy()
                dfp['dob_dt'] = pd.to_datetime(dfp['dob'], errors='coerce')
                dfp['age'] = ((pd.Timestamp(date.today()) - dfp['dob_dt']).dt.days // 365).fillna(0).astype(int)
                bins = pd.cut(dfp['age'], bins=[-1,0,1,5,12,18,40,65,200],
                              labels=["<1","1-5","6-12","13-18","19-40","41-65","66+"])
                age_count = bins.value_counts().sort_index()
                fig, ax = plt.subplots(figsize=(6,3))
                age_count.plot(kind='bar', color=ACCENT, ax=ax)
                ax.set_title("Patients by age group")
                ax.set_xlabel("")
                png = _save_png_from_matplotlib(fig)
                charts["Patients by age group"] = png
                plt.close(fig)
        except Exception:
            pass
        word_bytes = create_word_report(patients_df, staff_df, schedule_df, charts=charts if charts else None)
        st.download_button("Download Word report (with charts)", data=word_bytes, file_name="homecare_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.markdown("---")
    st.write("Backup: download underlying SQLite DB")
    try:
        with open(DB_PATH, "rb") as f:
            db_bytes = f.read()
            st.download_button("Download DB file", data=db_bytes, file_name=DB_PATH, mime="application/x-sqlite3")
    except Exception as e:
        st.error("Could not read DB file: " + str(e))

    render_footer(show_all_rights=True)

# ---------- Logout ----------
elif choice == "Logout":
    st.subheader("Logout")
    st.write(f"Logged in as **{st.session_state.user}** ({st.session_state.user_role})")
    if st.button("Logout"):
        logout_user()
        st.success("Logged out")
        st.experimental_rerun()

# End of file
